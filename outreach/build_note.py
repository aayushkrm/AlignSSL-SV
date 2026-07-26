"""Build the science-popular note as a .docx.

Structure follows the six-part skeleton from the curator's guide:
title/abstract, project, known results, our results, team process, plans.
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn

FONT = "Calibri"
BLUE = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x5A, 0x5F, 0x66)

doc = Document()

# --- global style ---------------------------------------------------------
st = doc.styles["Normal"]
st.font.name = FONT
st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
pf = st.paragraph_format
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pf.space_after = Pt(8)
pf.line_spacing = 1.15

for name, size, colour, bold in [("Heading 1", 16, BLUE, True),
                                 ("Heading 2", 13, BLUE, True),
                                 ("Title", 22, BLUE, True)]:
    s = doc.styles[name]
    s.font.name = FONT
    s.font.size = Pt(size)
    s.font.color.rgb = colour
    s.font.bold = bold

for s in doc.sections:
    s.left_margin = s.right_margin = Inches(1.0)
    s.top_margin = s.bottom_margin = Inches(0.9)


def h1(t):
    doc.add_heading(t, level=1)


def h2(t):
    doc.add_heading(t, level=2)


def p(t, style=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    par = doc.add_paragraph(t, style=style)
    par.alignment = align
    return par


def bullet(t):
    par = doc.add_paragraph(t, style="List Bullet")
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return par


def figure(path, caption, width=6.3):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_after = Pt(4)
    par.add_run().add_picture(path, width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(14)
    r = cap.add_run(caption)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = GREY


def equation(path, caption, width=None):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(6)
    par.paragraph_format.space_after = Pt(2)
    par.add_run().add_picture(path, width=Inches(width) if width else None)
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        r = cap.add_run(caption)
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = GREY


# =========================================================================
t = doc.add_paragraph("Teaching a machine to spot missing DNA — "
                      "and catching the test that was too easy",
                      style="Title")
t.alignment = WD_ALIGN_PARAGRAPH.LEFT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = sub.add_run("How learning without answers made our deletion caller work "
                "from a handful of labelled examples — and why the more "
                "interesting result was a negative one")
r.font.size = Pt(12)
r.font.italic = True
r.font.color.rgb = GREY

by = doc.add_paragraph()
by.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = by.add_run("Igor Novikov (curator) · Ayush Kumar (team lead) · "
               "Nicolas Baranov · Alisa Chernysheva · Baowen Zeng 曾宝文 · "
               "Haoxin Bai 白昊欣 · Zicong Guan 关紫聪")
r.font.size = Pt(10)
r.font.color.rgb = GREY

h1("In brief")
p("Every human genome is missing chunks of DNA that other genomes have. Some "
  "of those missing chunks cause disease. Finding them in raw sequencing data "
  "is a decade-old problem, and since 2019 the standard approach has been to "
  "draw the data as a colour picture and hand it to an image-recognition "
  "neural network. We asked two questions. First: can a network learn what "
  "normal sequencing data looks like without being told any answers, so that "
  "it then needs far fewer labelled examples? Second — and this is the "
  "question that changed the paper — is the benchmark everyone uses actually "
  "hard enough to tell good methods from bad ones?")
p("The answers were yes, and no. Pre-training without labels let our model "
  "reach an accuracy of 0.51 from just 210 labelled examples, where the same "
  "network trained the ordinary way managed 0.05. But twelve simple numbers "
  "you can compute with a pocket calculator, fed to a standard machine-"
  "learning method, beat every neural network on the same test. That is a "
  "negative result about the benchmark, not about the networks — and we "
  "think reporting it is more useful than another leaderboard entry.")

# ---------------------------------------------------------------- section 2
h1("The problem: reading what is not there")
p("A genome is three billion letters long. A sequencing machine does not read "
  "it end to end; it shreds the DNA into tens of millions of short fragments "
  "of roughly a hundred letters each, reads those, and leaves you to work out "
  "where each fragment came from. Software called an aligner does that "
  "matching, pinning every fragment — a “read” — onto a reference copy of the "
  "human genome.")
p("A deletion is a stretch of DNA, anywhere from fifty letters to hundreds of "
  "thousands, that a particular person simply does not have. Because the "
  "person does not have it, no read can come from it. The evidence is "
  "therefore indirect, and it arrives in three forms at once: a region where "
  "the pile of reads suddenly thins out or vanishes; pairs of reads that were "
  "cut from opposite ends of one fragment but land much further apart than "
  "they should, because the DNA between them is gone; and reads that straddle "
  "the boundary and get sliced in half by the aligner, matching cleanly on one "
  "side and not at all on the other.")
figure("note_fig1_schematic.png",
       "Figure 1. What a deletion does to sequencing data, and two ways of "
       "showing that evidence to a neural network. Panel a: reads thin out "
       "over the deleted stretch, paired reads land too far apart, and reads "
       "crossing the boundary are cut short. Panel b: the 2019 approach paints "
       "this into a colour image. Panel c: our approach keeps the numbers the "
       "aligner already computed, one layer per kind of evidence.")
p("Deletions matter clinically. They are enriched among disease-causing "
  "variants, and collectively they change more letters of a genome than the "
  "single-letter typos that get most of the attention. They are also the "
  "hardest common class of variation to call reliably: different tools "
  "disagree with each other on the same data, and their confidence scores are "
  "rarely trustworthy enough to act on.")

h2("The same problem, stated mathematically")
p("Take a window of the genome around a candidate site. Stack the reads that "
  "overlap it, one per row, and for each read record what the aligner already "
  "knows at each position: whether the base matches the reference, how "
  "confident the aligner is about where the read belongs, how far away its "
  "partner landed, whether it was cut short, which strand it came from, how "
  "deep the pile is. That gives a three-dimensional array of numbers — an "
  "alignment tensor:")
equation("eq1_tensor.png", "", width=2.3)
p("Eighteen kinds of evidence, sixty-four reads, two hundred and fifty-six "
  "positions. The task is then a function that maps this array to a single "
  "answer: is there a deletion here or not. Learning that function from "
  "examples is what the rest of this note is about.")

# ---------------------------------------------------------------- section 3
h1("What was already known, and what was wrong with it")
p("In 2019 a paper called DeepSV made an influential move. Rather than "
  "hand-engineering rules about read depth and insert size, it rendered the "
  "pile of reads as an ordinary RGB image — base identity, base quality and "
  "strand encoded as colours — and trained an image classifier on it. It "
  "worked, and it started a whole family of “pileup image” methods that are "
  "still the mainstream today.")
p("But that first move fixed three decisions that the field has largely "
  "inherited without re-examining:")
bullet("The picture is a human invention, not something the model learns. "
       "Whatever does not survive the translation into colours — the fine "
       "structure of insert sizes, the geometry of where reads got cut — is "
       "thrown away before the network ever sees it.")
bullet("Training needs a lot of labelled answers. Every new sequencing "
       "machine, coverage level, or population needs a fresh truth set, and "
       "reliable truth sets exist for only a handful of reference samples.")
bullet("The confidence numbers do not mean what they look like. A network "
       "that says “95 % sure” is not right 95 % of the time, which makes the "
       "output hard to threshold in a clinic.")
p("Meanwhile, machine learning as a field had developed a direct remedy for "
  "the second problem: self-supervised pre-training. Show a model enormous "
  "amounts of unlabelled data, hide part of it, and make the model fill in "
  "what was hidden. To do that well it has to learn the structure of the data "
  "— and it can then be fine-tuned on a small labelled set. This is how modern "
  "language models are built. In genomics it has been applied to the reference "
  "DNA sequence, but essentially not to the alignment evidence that actually "
  "tells you a deletion is there.")

# ---------------------------------------------------------------- section 4
h1("What we did and what we found")
h2("Learning without answers")
p("We built the alignment tensor described above and trained an encoder on "
  "120,000 windows taken from three people, using no deletion labels at all. "
  "The training signal has two parts. We hide a random 60 % of the entries in "
  "the tensor and ask the model to reconstruct them; and we show it two "
  "differently-corrupted views of the same window and require the two "
  "resulting summaries to agree, while a regulariser stops the model from "
  "cheating by making every summary identical:")
equation("eq2_loss.png", "", width=5.6)
p("Only after that do labels enter, and only to train a small classification "
  "head on top of the frozen-then-fine-tuned encoder. The whole point is to "
  "measure how few labels are enough.")

h2("The headline: a working detector from 210 examples")
p("We compared four things on data from the 1000 Genomes Project — a panel of "
  "six people spanning five continental ancestries, with a seventh held back "
  "entirely for testing. The comparison covers our tensor model with "
  "pre-training, the identical network trained from scratch, a re-"
  "implementation of the 2019 colour-image representation, and, as a sanity "
  "check, twelve hand-computed summary numbers fed to classical machine "
  "learning.")
figure("note_fig2_results.png",
       "Figure 2. Panel a: accuracy against the number of labelled examples "
       "used, for our network with and without self-supervised pre-training. "
       "The advantage is largest where labels are scarcest. Panel b: how well "
       "a single simple measurement, on its own and with no training at all, "
       "separates real deletions from decoys — on the original benchmark and "
       "on our repaired one.")
p("Panel a is the result we set out to get. With 210 labelled examples — one "
  "per cent of the training set — the pre-trained model reaches an accuracy "
  "of 0.51 while the same architecture from scratch reaches 0.05 — a tenfold "
  "difference in accuracy. By five per cent of labels the from-scratch model "
  "has caught up, and at full supervision it is in fact a hair ahead. "
  "Pre-training buys label efficiency, not a higher ceiling, and we say so.")

h2("The result we did not go looking for")
p("Then the sanity check misbehaved. Those twelve hand-computed numbers, fed "
  "to a gradient-boosted tree, scored 0.894 from the same 210 labels — better "
  "than every neural network in the comparison. Worse, a single one of them, "
  "the ratio of read depth in the middle of the window to read depth at its "
  "edges, separated deletions from non-deletions at an AUC of 0.955 with no "
  "training whatsoever.")
p("That is not a story about neural networks. It is a story about the test. "
  "The standard way to build such a benchmark is to take known deletions as "
  "positive examples and pick random genomic windows as negatives. Random "
  "windows are almost never anything like a deletion, so a coverage dip alone "
  "gives the answer away. Every method in the comparison, ours included, was "
  "largely being scored on how fast it rediscovers one obvious heuristic.")
p("So we rebuilt the benchmark. Instead of random negatives we selected "
  "decoys that a real caller would actually have to adjudicate — windows "
  "matched to the true deletions on their depth profile, so the easy shortcut "
  "is taken away. On this harder task the depth ratio drops from 0.955 to "
  "0.717 (Figure 2b): attenuated, not eliminated, and we report it that way.")
figure("note_fig3_standings.png",
       "Figure 3. Final accuracy of every approach once the easy shortcut has "
       "been removed. Bars show the mean over three independent runs; whiskers "
       "show the spread between runs.")
p("Every method falls on the harder benchmark, which is the point of a harder "
  "benchmark. Three things survive the change. The label-efficiency result not "
  "only survives but sharpens: with one per cent of labels the pre-trained "
  "model scores 0.352 and the from-scratch model scores exactly zero. The "
  "2019-style colour representation collapses to 0.28, well below everything "
  "else, which is evidence that keeping the aligner's numbers instead of "
  "painting them into pixels is worth doing. And the hand-crafted features "
  "still lead, at 0.79 against our 0.76.")
p("We could have quietly kept the first benchmark, reported the low-label "
  "result, and stopped. The comparison would have looked better. We "
  "think the field is better served by the honest version: pileup-style "
  "benchmarks built this way cannot support claims about how a caller will "
  "behave in practice, and fixing the negative sampling helps but does not "
  "finish the job.")

# ---------------------------------------------------------------- section 5
h1("How the team worked")
p("Seven of us, in three countries and several time zones, working on a "
  "shared university cluster with a hard limit of ten CPUs per user and four "
  "GPUs shared with everyone else. Almost every practical decision in the "
  "project was shaped by that constraint rather than by the science.")
p("The work split naturally into streams that could run in parallel: getting "
  "sequencing data down from public archives and indexed; turning alignments "
  "into tensors at scale; the self-supervised pre-training runs; the "
  "supervised evaluation and baselines; benchmark design and the statistics; "
  "and writing. People moved between streams as the bottleneck moved.")
p("Three habits did most of the work. Everything lived in one public Git "
  "repository from the first week, so there was never a private version of "
  "the truth. Every number that appears in a figure or in the manuscript is "
  "read at build time from a results table that a job on the cluster wrote — "
  "nobody types a number by hand, and an automated check refuses to build the "
  "paper if a quoted value and its table disagree. And we ran deliberate "
  "adversarial review passes on our own work, where one person's job was to "
  "attack the result rather than defend it.")
p("That last habit is what produced the interesting half of this project. The "
  "hand-crafted-feature control was added as a formality, to have something "
  "unflattering to compare against. It won. Taking that seriously, rather "
  "than filing it as an anomaly, is the reason there is a negative result to "
  "report.")
p("What we learned along the way was mostly unglamorous and mostly "
  "transferable: that a shared filesystem can disappear overnight and your "
  "recovery plan is only as good as what you wrote down; that a "
  "sixty-six-gigabyte array read from network storage will starve a GPU until "
  "you stage it into local memory; that a test which imports a library the "
  "cluster does not have will fail two hours into a job instead of two seconds "
  "into it, unless you write it not to; and that the difference between a "
  "result and a publishable result is mostly the controls you ran.")

# ---------------------------------------------------------------- section 6
h1("What is left")
p("Several things did not get done. We had planned to validate against the "
  "Genome in a Bottle reference sample, the gold standard in this field, and "
  "that is still outstanding. Partway through, the shared data workspace "
  "holding our sequencing files was lost, and only two of the seven "
  "alignments survived; the repaired benchmark therefore rests on a single "
  "sample rather than the full panel, which is a real limitation and is "
  "labelled as one in the paper. And we only ever attacked deletions. "
  "Insertions, inversions and duplications are all harder and all left open.")
p("The near-term plan is to close the gold-standard validation, put the "
  "repaired benchmark on the full panel once the data are restaged, and add "
  "a head that predicts where the deletion starts and ends rather than only "
  "whether one is present. The longer-term question is the one this project "
  "backed into: if the standard benchmarks in a field cannot distinguish a "
  "twelve-number heuristic from a deep network, then building better "
  "benchmarks is worth more than building better networks, and that is where "
  "we would rather spend the next year.")
p("The code, the extraction pipeline, every results table and every figure "
  "are public at github.com/aayushkrm/AlignSSL-SV under an MIT licence.")

doc.save("AlignSSL_SV_popular_note.docx")
print("saved")
